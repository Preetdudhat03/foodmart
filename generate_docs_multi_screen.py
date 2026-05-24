import os
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from playwright.sync_api import sync_playwright

experiments = [
    {
        "id": "2",
        "aim": "multicolumn layout using bootstrap grid system responsive breakpoint for mobile, tablet and the desktop",
        "code": """// File: shop.php
<div class="product-grid row row-cols-1 row-cols-sm-2 row-cols-md-3 row-cols-lg-3 row-cols-xl-4">
  <?php if ($result->num_rows > 0): ?>
    <?php while($row = $result->fetch_assoc()): ?>
    <div class="col">
      <div class="product-item">
         <!-- Product Display -->
         <img src="<?php echo htmlspecialchars($row['image']); ?>">
         <h3><?php echo htmlspecialchars($row['title']); ?></h3>
         <span class="price"><?php echo htmlspecialchars($row['price']); ?></span>
      </div>
    </div>
    <?php endwhile; ?>
  <?php endif; ?>
</div>""",
        "url": "http://localhost/foodmart/shop.php",
        "steps": [
            {
                "name": "Desktop Layout",
                "action": lambda page: page.set_viewport_size({"width": 1280, "height": 800})
            },
            {
                "name": "Mobile Layout (Grid stacked)",
                "action": lambda page: page.set_viewport_size({"width": 400, "height": 800})
            }
        ]
    },
    {
        "id": "3",
        "aim": "form control and validation",
        "code": """// File: account.php
<form action="actions/change_password.php" method="POST">
  <div class="mb-3">
    <label for="current_password" class="form-label">Current Password</label>
    <input type="password" class="form-control" id="current_password" name="current_password" required>
  </div>
  <div class="mb-3">
    <label for="new_password" class="form-label">New Password</label>
    <input type="password" class="form-control" id="new_password" name="new_password" required>
  </div>
  <div class="d-grid gap-2">
    <button type="submit" class="btn btn-danger">Update Password</button>
  </div>
</form>""",
        "url": "http://localhost/foodmart/index.php",
        "steps": [
            {
                "name": "Login Form View",
                "action": lambda page: (
                    page.evaluate("() => { const el = document.querySelector('a[href=\"#loginModal\"]'); if(el) el.click(); }"),
                    page.wait_for_timeout(1000)
                )
            },
            {
                "name": "Validation Triggered",
                "action": lambda page: (
                    page.evaluate("() => { const el = document.querySelector('#loginForm button[type=\"submit\"]'); if(el) el.click(); }"),
                    page.wait_for_timeout(500)
                )
            }
        ]
    },
    {
        "id": "4",
        "aim": "simple DOM manipulation",
        "code": """// File: js/cart.js
// Update Cart Badge dynamically
function updateCartBadge(count) {
    $('#offcanvasCart .badge').text(count);
}
// jQuery manipulates the DOM elements to reflect the new count.
""",
        "url": "http://localhost/foodmart/shop.php",
        "steps": [
            {
                "name": "Initial State",
                "action": lambda page: page.wait_for_timeout(500)
            },
            {
                "name": "DOM Manipulated (Item added to cart)",
                "action": lambda page: (
                    page.evaluate("() => { const btns = document.querySelectorAll('.btn-add-to-cart'); if(btns.length > 0) btns[0].click(); }"),
                    page.wait_for_timeout(1500)
                )
            }
        ]
    },
    {
        "id": "5",
        "aim": "form validation using string method and type conversion",
        "code": """// File: js/auth.js
// Register Form Submission (AJAX JSON validation handling)
$('#registerForm').on('submit', function (e) {
    e.preventDefault();
    $.ajax({
        type: 'POST',
        url: 'actions/register.php',
        data: $(this).serialize(),
        dataType: 'json',
        success: function (response) {
            if (!response.success) {
                alert(response.message);
            }
        }
    });
});""",
        "url": "http://localhost/foodmart/index.php",
        "steps": [
            {
                "name": "Registration Form",
                "action": lambda page: (
                    page.evaluate("() => { document.querySelector('#offcanvasUser').classList.add('show'); }"),
                    page.wait_for_timeout(500),
                    page.evaluate("() => { const el = document.querySelector('#pills-register-tab'); if(el) el.click(); }"),
                    page.wait_for_timeout(500)
                )
            }
        ]
    },
    {
        "id": "6",
        "aim": "demonstrate the functions and variable hoisting",
        "code": """// File: js/cart.js
$(document).ready(function () {
    // Hoisting in action - calling function before declaration
    updateCartBadge(0); 

    function updateCartBadge(count) {
        $('#offcanvasCart .badge').text(count);
    }
});""",
        "url": "http://localhost/foodmart/index.php",
        "steps": [
            {
                "name": "Cart Icon initialized via Hoisted function",
                "action": lambda page: page.wait_for_timeout(500)
            }
        ]
    },
    {
        "id": "7",
        "aim": "jquery event listener and animation",
        "code": """// File: js/cart.js
// jQuery event listener and icon animation
$(document).on('click', '.btn-add-to-cart', function (e) {
    e.preventDefault();
    var $btn = $(this);
    var originalIcon = $btn.html();
    $btn.html('<i class="bi bi-check-lg"></i>'); // UI feedback animation

    setTimeout(function () {
        $btn.html(originalIcon);
    }, 1500);
});""",
        "url": "http://localhost/foodmart/shop.php",
        "steps": [
            {
                "name": "Animation triggered on click",
                "action": lambda page: (
                    page.evaluate("() => { const btns = document.querySelectorAll('.btn-add-to-cart'); if(btns.length > 1) btns[1].click(); }"),
                    page.wait_for_timeout(200) # Quick snapshot to catch the checkmark
                )
            }
        ]
    },
    {
        "id": "8",
        "aim": "fetches data asynchronously using jquery",
        "code": """// File: js/cart.js
$.ajax({
    url: 'actions/cart_actions.php',
    type: 'POST',
    dataType: 'json',
    data: { action: 'add', id: id, quantity: quantity },
    success: function (response) {
        if (response.success) {
            updateCartBadge(response.cart_count);
        }
    }
});""",
        "url": "http://localhost/foodmart/shop.php",
        "steps": [
            {
                "name": "Sending Async Request",
                "action": lambda page: (
                    page.evaluate("() => { const btns = document.querySelectorAll('.btn-add-to-cart'); if(btns.length > 0) btns[0].click(); }"),
                    page.wait_for_timeout(500)
                )
            },
            {
                "name": "Async Result Displayed",
                "action": lambda page: (
                    page.evaluate("() => { document.querySelector('#offcanvasCart').classList.add('show'); }"),
                    page.wait_for_timeout(1000)
                )
            }
        ]
    },
    {
        "id": "9",
        "aim": "php scripts that perform sorting searching and merging",
        "code": """// File: shop.php
$category_filter = isset($_GET['category']) ? $_GET['category'] : 'all';
$sql = "SELECT * FROM products";
if ($category_filter != 'all') {
    $sql .= " WHERE category = ?";
    $params[] = $category_filter;
}
$stmt = $conn->prepare($sql);
$stmt->execute();
$result = $stmt->get_result();""",
        "url": "http://localhost/foodmart/shop.php",
        "steps": [
            {
                "name": "All Products",
                "action": lambda page: page.wait_for_timeout(500)
            },
            {
                "name": "Filtered Products",
                "action": lambda page: (
                    page.goto("http://localhost/foodmart/shop.php?category=Beverages"),
                    page.wait_for_load_state("networkidle")
                )
            }
        ]
    },
    {
        "id": "10",
        "aim": "laravel based crud using route controller include the user friendly error manage and bootstrap styling (Mapped to raw PHP Cart CRUD)",
        "code": """// File: cart.php & actions/cart_actions.php
// CREATE (Add) -> actions/cart_actions.php
$_SESSION['cart'][$id] = ['name' => $name, 'price' => $price, 'quantity' => $quantity];

// READ (Display) -> cart.php
<?php foreach ($_SESSION['cart'] as $id => $item): ?>
    <td><?php echo htmlspecialchars($item['name']); ?></td>
<?php endforeach; ?>

// DELETE (Remove) -> actions/cart_actions.php
unset($_SESSION['cart'][$id]);
""",
        "url": "http://localhost/foodmart/cart.php",
        "steps": [
            {
                "name": "Empty Cart (Initial State)",
                "action": lambda page: page.wait_for_timeout(500)
            },
            {
                "name": "Cart with Items (Create/Read)",
                "action": lambda page: (
                    page.goto("http://localhost/foodmart/shop.php"),
                    page.evaluate("() => { const btns = document.querySelectorAll('.btn-add-to-cart'); if(btns.length > 0) btns[0].click(); }"),
                    page.wait_for_timeout(1000),
                    page.goto("http://localhost/foodmart/cart.php"),
                    page.wait_for_load_state("networkidle")
                )
            }
        ]
    }
]

os.makedirs('c:/xampp/htdocs/FoodMart/exp_output', exist_ok=True)

print("Taking screenshots of FoodMart app...")
with sync_playwright() as p:
    browser = p.chromium.launch()
    page = browser.new_page()
    
    for exp in experiments:
        print(f"Processing Exp {exp['id']}...")
        try:
            page.set_viewport_size({"width": 1280, "height": 800})
            page.goto(exp['url'])
            page.wait_for_load_state("networkidle")
            
            exp['images'] = []
            
            for idx, step in enumerate(exp['steps']):
                if step['action']:
                    try:
                        step['action'](page)
                    except Exception as e:
                        print(f"Action failed for {exp['id']} step {idx}: {e}")
                
                screenshot_path = f"c:/xampp/htdocs/FoodMart/exp_output/screenshot_{exp['id']}_{idx}.png"
                page.screenshot(path=screenshot_path)
                exp['images'].append({
                    "path": screenshot_path,
                    "name": step['name']
                })
        except Exception as e:
            print(f"Failed {exp['id']}: {e}")

    browser.close()

print("Generating Documents...")
combined_doc = Document()
combined_doc.add_heading("FoodMart Code Experiments", level=1)

for idx, exp in enumerate(experiments):
    doc = Document()
    
    # Configure document to maximize space
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    title = doc.add_heading(f"Exp - {exp['id']}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading("Aim", level=2)
    doc.add_paragraph(exp['aim'])
    
    doc.add_heading("Code (From FoodMart Project)", level=2)
    code_paragraph = doc.add_paragraph()
    code_run = code_paragraph.add_run(exp['code'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)
    
    if 'images' in exp:
        for img in exp['images']:
            if os.path.exists(img['path']):
                p = doc.add_paragraph()
                r = p.add_run()
                r.add_picture(img['path'], width=Inches(4.5))
                p_text = doc.add_paragraph(img['name'])
                p_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.save(f"c:/xampp/htdocs/FoodMart/exp_output/Experiment_{exp['id']}.docx")
    
    # Combined Document
    if idx > 0:
        combined_doc.add_page_break()
        
    title = combined_doc.add_heading(f"Exp - {exp['id']}", level=2)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    combined_doc.add_heading("Aim", level=3)
    combined_doc.add_paragraph(exp['aim'])
    
    combined_doc.add_heading("Code (From FoodMart Project)", level=3)
    code_paragraph = combined_doc.add_paragraph()
    code_run = code_paragraph.add_run(exp['code'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)
    
    if 'images' in exp:
        for img in exp['images']:
            if os.path.exists(img['path']):
                p = combined_doc.add_paragraph()
                r = p.add_run()
                r.add_picture(img['path'], width=Inches(4.5))
                p_text = combined_doc.add_paragraph(img['name'])
                p_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

sections = combined_doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
combined_doc.save("c:/xampp/htdocs/FoodMart/exp_output/FoodMart_Experiments.docx")
print("Done!")
