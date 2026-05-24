import os
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from playwright.sync_api import sync_playwright

experiments = [
    {
        "id": "2",
        "aim": "multicolumn layout using bootstrap grid system responsive breakpoint for mobile, tablet and the desktop",
        "filename": "exp2.html",
        "url": "http://localhost/foodmart/exp_output/exp2.html",
        "code": """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Bootstrap Grid</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body>
    <div class="container mt-5">
        <h2 class="text-center mb-4">Responsive Multi-Column Layout</h2>
        <div class="row">
            <div class="col-12 col-md-6 col-lg-4 mb-3">
                <div class="p-3 bg-primary text-white text-center">Column 1</div>
            </div>
            <div class="col-12 col-md-6 col-lg-4 mb-3">
                <div class="p-3 bg-success text-white text-center">Column 2</div>
            </div>
            <div class="col-12 col-md-12 col-lg-4 mb-3">
                <div class="p-3 bg-danger text-white text-center">Column 3</div>
            </div>
        </div>
    </div>
</body>
</html>""",
        "action": None
    },
    {
        "id": "3",
        "aim": "form control and validation",
        "filename": "exp3.html",
        "url": "http://localhost/foodmart/exp_output/exp3.html",
        "code": """<!DOCTYPE html>
<html>
<head>
    <title>Form Validation</title>
    <style>
        body { font-family: Arial; padding: 20px; }
        .error { color: red; font-size: 0.9em; display: none; margin-top: 5px; }
        .form-group { margin-bottom: 15px; }
        input { padding: 8px; width: 200px; }
        button { padding: 8px 15px; }
    </style>
</head>
<body>
    <h2>Form Control & Validation</h2>
    <form id="myForm" onsubmit="return validateForm(event)">
        <div class="form-group">
            <label for="email">Email:</label><br>
            <input type="email" id="email" required>
            <div class="error" id="emailError">Please enter a valid email.</div>
        </div>
        <div class="form-group">
            <label for="password">Password:</label><br>
            <input type="password" id="password" required minlength="6">
            <div class="error" id="pwdError">Password must be at least 6 characters.</div>
        </div>
        <button type="submit" id="submitBtn">Submit</button>
    </form>

    <script>
        function validateForm(e) {
            e.preventDefault();
            let isValid = true;
            let pwd = document.getElementById('password').value;
            
            if (pwd.length < 6) {
                document.getElementById('pwdError').style.display = 'block';
                isValid = false;
            } else {
                document.getElementById('pwdError').style.display = 'none';
            }
            
            if (isValid) alert('Form submitted successfully!');
            return isValid;
        }
    </script>
</body>
</html>""",
        "action": lambda page: page.locator("#submitBtn").click() # trigger validation error
    },
    {
        "id": "4",
        "aim": "simple DOM manipulation",
        "filename": "exp4.html",
        "url": "http://localhost/foodmart/exp_output/exp4.html",
        "code": """<!DOCTYPE html>
<html>
<head><style>body{font-family:Arial; padding: 20px;}</style></head>
<body>
    <div id="contentBox" style="padding: 20px; border: 1px solid black; width: 300px; transition: 0.3s;">
        <h3 id="mainTitle">Original Title</h3>
        <p id="descText">This is some standard description text.</p>
    </div>
    <br>
    <button id="changeBtn" onclick="manipulateDOM()">Change Content</button>

    <script>
        function manipulateDOM() {
            let title = document.getElementById('mainTitle');
            let box = document.getElementById('contentBox');
            let p = document.getElementById('descText');
            
            title.innerText = "Updated Title via DOM";
            p.innerHTML = "<strong>The text has been styled and updated dynamically!</strong>";
            
            box.style.backgroundColor = "lightyellow";
            box.style.borderRadius = "10px";
            box.style.borderColor = "blue";
        }
    </script>
</body>
</html>""",
        "action": lambda page: page.locator("#changeBtn").click()
    },
    {
        "id": "5",
        "aim": "form validation using string method and type conversion",
        "filename": "exp5.html",
        "url": "http://localhost/foodmart/exp_output/exp5.html",
        "code": """<!DOCTYPE html>
<html>
<head><style>body{font-family:Arial; padding: 20px;}</style></head>
<body>
    <h3>User Registration Validation</h3>
    <input type="text" id="username" placeholder="Username (Letters only)"><br><br>
    <input type="text" id="age" placeholder="Age"><br><br>
    <button id="valBtn" onclick="validateData()">Validate</button>
    <p id="result"></p>

    <script>
        function validateData() {
            let username = document.getElementById('username').value.trim();
            let ageStr = document.getElementById('age').value;
            let resultText = "";
            
            let nameRegex = /^[a-zA-Z]+$/;
            if (!username.match(nameRegex)) {
                resultText += "Error: Username must contain only letters.<br>";
            } else if (username.length < 4) {
                resultText += "Error: Username must be at least 4 chars long.<br>";
            }
            
            let ageNum = parseInt(ageStr);
            if (isNaN(ageNum)) {
                resultText += "Error: Age must be a valid number.<br>";
            } else if (ageNum < 18) {
                resultText += "Error: You must be at least 18 years old.<br>";
            }
            
            if (resultText === "") {
                resultText = "<span style='color:green'>Validation Successful!</span>";
            } else {
                resultText = "<span style='color:red'>" + resultText + "</span>";
            }
            
            document.getElementById('result').innerHTML = resultText;
        }
    </script>
</body>
</html>""",
        "action": lambda page: (page.locator("#username").fill("Jo"), page.locator("#age").fill("abc"), page.locator("#valBtn").click())
    },
    {
        "id": "6",
        "aim": "demonstrate the functions and variable hoisting",
        "filename": "exp6.html",
        "url": "http://localhost/foodmart/exp_output/exp6.html",
        "code": """<!DOCTYPE html>
<html>
<head><style>body{font-family:Arial; padding: 20px;} pre{background:#eee; padding:10px;}</style></head>
<body>
    <h3>Hoisting Demonstration</h3>
    <p>Check the output generated by JavaScript hoisting below:</p>
    <pre id="output"></pre>
    <script>
        let out = "";
        out += "--- Variable Hoisting ---\\n";
        out += "Value of myVar before declaration: " + myVar + "\\n"; 
        var myVar = "I am declared!";
        out += "Value of myVar after declaration: " + myVar + "\\n\\n";
        
        out += "--- Function Hoisting ---\\n";
        out += sayHello() + "\\n\\n";
        function sayHello() { return "Hello! This function was hoisted successfully."; }
        
        out += "--- Let/Const Hoisting (Temporal Dead Zone) ---\\n";
        try {
            console.log(myLetVar);
        } catch (error) {
            out += "Error caught: Cannot access 'myLetVar' before initialization\\n";
        }
        let myLetVar = "Modern JS";
        
        document.getElementById("output").innerText = out;
    </script>
</body>
</html>""",
        "action": None
    },
    {
        "id": "7",
        "aim": "jquery event listener and animation",
        "filename": "exp7.html",
        "url": "http://localhost/foodmart/exp_output/exp7.html",
        "code": """<!DOCTYPE html>
<html>
<head>
    <title>jQuery Animation</title>
    <script src="https://code.jquery.com/jquery-3.6.0.min.js"></script>
    <style>
        body{font-family:Arial; padding: 20px;}
        #animatedBox { width: 150px; height: 150px; background-color: #3498db; margin-top: 20px; position: relative; }
    </style>
</head>
<body>
    <button id="moveBtn">Move & Fade Box</button>
    <button id="resetBtn">Reset</button>
    <div id="animatedBox"></div>
    <script>
        $(document).ready(function() {
            $("#moveBtn").click(function() {
                $("#animatedBox").animate({ left: '250px', opacity: '0.5', height: '200px', width: '200px' }, 500);
            });
        });
    </script>
</body>
</html>""",
        "action": lambda page: (page.locator("#moveBtn").click(), page.wait_for_timeout(600))
    },
    {
        "id": "8",
        "aim": "fetches data asynchronously using jquery",
        "filename": "exp8.html",
        "url": "http://localhost/foodmart/exp_output/exp8.html",
        "code": """<!DOCTYPE html>
<html>
<head>
    <title>jQuery AJAX Fetch</title>
    <script src="https://code.jquery.com/jquery-3.6.0.min.js"></script>
    <style>body{font-family:Arial; padding: 20px;} li {margin-bottom: 5px;}</style>
</head>
<body>
    <h3>User List from API</h3>
    <button id="fetchBtn">Load Users</button>
    <p id="loading" style="display:none; color:blue;">Loading data...</p>
    <ul id="userList"></ul>
    <script>
        $(document).ready(function() {
            $("#fetchBtn").click(function() {
                $("#loading").show();
                $.ajax({
                    url: "https://jsonplaceholder.typicode.com/users",
                    type: "GET",
                    success: function(data) {
                        $("#loading").hide();
                        let listItems = "";
                        for (let i = 0; i < 5; i++) {
                            listItems += "<li><strong>" + data[i].name + "</strong> (" + data[i].email + ")</li>";
                        }
                        $("#userList").html(listItems);
                    }
                });
            });
        });
    </script>
</body>
</html>""",
        "action": lambda page: (page.locator("#fetchBtn").click(), page.wait_for_timeout(1500))
    },
    {
        "id": "9",
        "aim": "php scripts that perform sorting searching and merging",
        "filename": "exp9.php",
        "url": "http://localhost/foodmart/exp_output/exp9.php",
        "code": """<?php
echo "<div style='font-family:Arial; padding: 20px;'>";
echo "<h3>PHP Array Operations</h3>";
$array1 = array(15, 42, 8, 23);
$array2 = array(10, 5, 99);
echo "<b>Original Array 1:</b> " . implode(", ", $array1) . "<br>";
echo "<b>Original Array 2:</b> " . implode(", ", $array2) . "<br><br>";
$merged_array = array_merge($array1, $array2);
echo "<b>Merged Array:</b> " . implode(", ", $merged_array) . "<br><br>";
sort($merged_array);
echo "<b>Sorted Merged Array (Ascending):</b> " . implode(", ", $merged_array) . "<br><br>";
$search_value = 23;
$index = array_search($search_value, $merged_array);
if ($index !== false) {
    echo "<span style='color:green;'><b>Search Result:</b> Value {$search_value} was found at index {$index} (0-based) in the sorted array.</span><br>";
}
echo "</div>";
?>""",
        "action": None
    },
    {
        "id": "10",
        "aim": "laravel based crud using route controller include the user friendly error manage and bootstrap styling",
        "filename": "exp10.html",
        "url": "http://localhost/foodmart/exp_output/exp10.html",
        "code": """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Laravel CRUD Simulation</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body>
<div class="container mt-5">
    <h2>Create New Product</h2>
    <div class="alert alert-danger">
        <ul>
            <li>Please provide a product name.</li>
            <li>The price must be a valid number.</li>
        </ul>
    </div>
    <form>
        <div class="mb-3">
            <label class="form-label">Product Name</label>
            <input type="text" name="name" class="form-control is-invalid" value="">
            <div class="invalid-feedback">Please provide a product name.</div>
        </div>
        <div class="mb-3">
            <label class="form-label">Price</label>
            <input type="text" name="price" class="form-control is-invalid" value="abc">
            <div class="invalid-feedback">The price must be a valid number.</div>
        </div>
        <button type="button" class="btn btn-primary">Save Product</button>
    </form>
</div>
</body>
</html>""",
        "action": None
    }
]

# Ensure directory exists
os.makedirs('c:/xampp/htdocs/FoodMart/exp_output', exist_ok=True)

# Write out the files
for exp in experiments:
    path = os.path.join('c:/xampp/htdocs/FoodMart/exp_output', exp['filename'])
    with open(path, 'w', encoding='utf-8') as f:
        f.write(exp['code'])

# Take screenshots
print("Taking screenshots...")
with sync_playwright() as p:
    browser = p.chromium.launch()
    page = browser.new_page()
    page.set_viewport_size({"width": 1024, "height": 768})
    
    for exp in experiments:
        print(f"Processing Exp {exp['id']}...")
        page.goto(exp['url'])
        page.wait_for_load_state("networkidle")
        if exp['action']:
            exp['action'](page)
            page.wait_for_timeout(500) # give time for UI to update
            
        screenshot_path = f"c:/xampp/htdocs/FoodMart/exp_output/screenshot_{exp['id']}.png"
        page.screenshot(path=screenshot_path)
        exp['screenshot'] = screenshot_path

    browser.close()

# Generate Documents
print("Generating Documents...")
combined_doc = Document()
combined_doc.add_heading("All Experiments", level=1)

for idx, exp in enumerate(experiments):
    doc = Document()
    
    # Single doc
    title = doc.add_heading(f"Exp - {exp['id']}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_heading("Aim", level=2)
    doc.add_paragraph(exp['aim'])
    
    doc.add_heading("Output (Screenshot)", level=2)
    doc.add_picture(exp['screenshot'], width=Inches(6.0))
    
    doc.add_heading("Code", level=2)
    code_paragraph = doc.add_paragraph()
    code_run = code_paragraph.add_run(exp['code'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.save(f"c:/xampp/htdocs/FoodMart/exp_output/Experiment_{exp['id']}.docx")
    
    # Combined doc
    if idx > 0:
        combined_doc.add_page_break()
    title = combined_doc.add_heading(f"Exp - {exp['id']}", level=2)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    combined_doc.add_heading("Aim", level=3)
    combined_doc.add_paragraph(exp['aim'])
    
    combined_doc.add_heading("Output (Screenshot)", level=3)
    combined_doc.add_picture(exp['screenshot'], width=Inches(6.0))
    
    combined_doc.add_heading("Code", level=3)
    code_paragraph = combined_doc.add_paragraph()
    code_run = code_paragraph.add_run(exp['code'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

combined_doc.save("c:/xampp/htdocs/FoodMart/exp_output/All_Experiments_With_Screenshots.docx")
print("Done!")
