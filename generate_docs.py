import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

experiments = [
    {
        "id": "2",
        "aim": "multicolumn layout using bootstrap grid system responsive breakpoint for mobile, tablet and the desktop",
        "code": """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Bootstrap Grid</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body>
    <div class="container mt-5">
        <h2 class="text-center mb-4">Responsive Multi-Column Layout</h2>
        <div class="row">
            <!-- Full width on mobile (col-12), half on tablet (col-md-6), third on desktop (col-lg-4) -->
            <div class="col-12 col-md-6 col-lg-4 mb-3">
                <div class="p-3 bg-primary text-white text-center">Column 1</div>
            </div>
            <div class="col-12 col-md-6 col-lg-4 mb-3">
                <div class="p-3 bg-success text-white text-center">Column 2</div>
            </div>
            <div class="col-12 col-md-12 col-lg-4 mb-3">
                <div class="p-3 bg-danger text-white text-center">Column 3</div>
            </div>
        </div>
    </div>
</body>
</html>""",
        "output": "The output renders a webpage with three colorful blocks.\n\nOn Desktop (Large Screen): The three blocks are displayed side-by-side in a single row.\nOn Tablet (Medium Screen): The first two blocks are displayed side-by-side (50% width), and the third block drops to the next row taking full width.\nOn Mobile (Small Screen): All three blocks stack vertically on top of each other, taking 100% width."
    },
    {
        "id": "3",
        "aim": "form control and validation",
        "code": """<!DOCTYPE html>
<html>
<head>
    <title>Form Validation</title>
    <style>
        .error { color: red; font-size: 0.9em; display: none; }
    </style>
</head>
<body>
    <form id="myForm" onsubmit="return validateForm()">
        <label for="email">Email:</label>
        <input type="email" id="email" required>
        <span class="error" id="emailError">Please enter a valid email.</span>
        <br><br>
        
        <label for="password">Password:</label>
        <input type="password" id="password" required minlength="6">
        <span class="error" id="pwdError">Password must be at least 6 characters.</span>
        <br><br>
        
        <button type="submit">Submit</button>
    </form>

    <script>
        function validateForm() {
            let isValid = true;
            let pwd = document.getElementById('password').value;
            
            if (pwd.length < 6) {
                document.getElementById('pwdError').style.display = 'block';
                isValid = false;
            } else {
                document.getElementById('pwdError').style.display = 'none';
            }
            
            if (isValid) alert('Form submitted successfully!');
            return isValid;
        }
    </script>
</body>
</html>""",
        "output": "Test Case 1 (Invalid Input):\nUser enters password '123' and clicks Submit.\nOutput: The form prevents submission, and the red error text 'Password must be at least 6 characters.' is displayed.\n\nTest Case 2 (Valid Input):\nUser enters valid email and password 'secret123' and clicks Submit.\nOutput: An alert pops up saying 'Form submitted successfully!'."
    },
    {
        "id": "4",
        "aim": "simple DOM manipulation",
        "code": """<!DOCTYPE html>
<html>
<body>
    <div id="contentBox" style="padding: 20px; border: 1px solid black; width: 300px;">
        <h3 id="mainTitle">Original Title</h3>
        <p id="descText">This is some standard description text.</p>
    </div>
    <br>
    <button onclick="manipulateDOM()">Change Content</button>

    <script>
        function manipulateDOM() {
            // Selecting elements
            let title = document.getElementById('mainTitle');
            let box = document.getElementById('contentBox');
            let p = document.getElementById('descText');
            
            // Manipulating text and HTML
            title.innerText = "Updated Title via DOM";
            p.innerHTML = "<strong>The text has been styled and updated dynamically!</strong>";
            
            // Manipulating styles
            box.style.backgroundColor = "lightyellow";
            box.style.borderRadius = "10px";
            box.style.borderColor = "blue";
        }
    </script>
</body>
</html>""",
        "output": "Initial State:\nA white box with a black border containing 'Original Title' and 'This is some standard description text.'\n\nAfter clicking 'Change Content' button:\n- The title changes to 'Updated Title via DOM'.\n- The paragraph text changes to bolded 'The text has been styled and updated dynamically!'.\n- The box's background color changes to lightyellow, border becomes blue, and the corners become rounded."
    },
    {
        "id": "5",
        "aim": "form validation using string method and type conversion",
        "code": """<!DOCTYPE html>
<html>
<body>
    <h3>User Registration Validation</h3>
    <input type="text" id="username" placeholder="Username (Letters only)"><br><br>
    <input type="text" id="age" placeholder="Age"><br><br>
    <button onclick="validateData()">Validate</button>
    <p id="result"></p>

    <script>
        function validateData() {
            let username = document.getElementById('username').value.trim();
            let ageStr = document.getElementById('age').value;
            let resultText = "";
            
            // String method validation
            let nameRegex = /^[a-zA-Z]+$/;
            if (!username.match(nameRegex)) {
                resultText += "Error: Username must contain only letters.<br>";
            } else if (username.length < 4) {
                resultText += "Error: Username must be at least 4 chars long.<br>";
            }
            
            // Type conversion validation
            let ageNum = parseInt(ageStr);
            if (isNaN(ageNum)) {
                resultText += "Error: Age must be a valid number.<br>";
            } else if (ageNum < 18) {
                resultText += "Error: You must be at least 18 years old.<br>";
            }
            
            if (resultText === "") {
                resultText = "<span style='color:green'>Validation Successful!</span>";
            } else {
                resultText = "<span style='color:red'>" + resultText + "</span>";
            }
            
            document.getElementById('result').innerHTML = resultText;
        }
    </script>
</body>
</html>""",
        "output": "Test Case 1:\nUsername: 'John12', Age: 'abc'\nOutput: \nError: Username must contain only letters.\nError: Age must be a valid number.\n\nTest Case 2:\nUsername: 'Bob', Age: '15'\nOutput:\nError: Username must be at least 4 chars long.\nError: You must be at least 18 years old.\n\nTest Case 3:\nUsername: 'Alice', Age: '22'\nOutput: Validation Successful!"
    },
    {
        "id": "6",
        "aim": "demonstrate the functions and variable hoisting",
        "code": """// Hoisting Demonstration Script

console.log("--- Variable Hoisting ---");
// Trying to access 'myVar' before initialization
console.log("Value of myVar before declaration: " + myVar); 

var myVar = "I am declared!";
console.log("Value of myVar after declaration: " + myVar);

console.log("\\n--- Function Hoisting ---");
// Calling function before it's declared in the code
sayHello();

// Function declaration
function sayHello() {
    console.log("Hello! This function was hoisted successfully.");
}

console.log("\\n--- Let/Const Hoisting (Temporal Dead Zone) ---");
try {
    console.log(myLetVar);
} catch (error) {
    console.log("Error caught: Cannot access 'myLetVar' before initialization");
}
let myLetVar = "Modern JS";""",
        "output": """--- Variable Hoisting ---
Value of myVar before declaration: undefined
Value of myVar after declaration: I am declared!

--- Function Hoisting ---
Hello! This function was hoisted successfully.

--- Let/Const Hoisting (Temporal Dead Zone) ---
Error caught: Cannot access 'myLetVar' before initialization"""
    },
    {
        "id": "7",
        "aim": "jquery event listener and animation",
        "code": """<!DOCTYPE html>
<html>
<head>
    <title>jQuery Animation</title>
    <!-- Include jQuery -->
    <script src="https://code.jquery.com/jquery-3.6.0.min.js"></script>
    <style>
        #animatedBox {
            width: 150px;
            height: 150px;
            background-color: #3498db;
            margin-top: 20px;
            position: relative;
        }
    </style>
</head>
<body>

    <button id="moveBtn">Move & Fade Box</button>
    <button id="resetBtn">Reset</button>

    <div id="animatedBox"></div>

    <script>
        $(document).ready(function() {
            // Event Listener for Animation Button
            $("#moveBtn").click(function() {
                $("#animatedBox").animate({
                    left: '250px',
                    opacity: '0.5',
                    height: '200px',
                    width: '200px'
                }, 1000); // 1000ms (1 second) duration
            });

            // Event Listener to reset
            $("#resetBtn").click(function() {
                $("#animatedBox").animate({
                    left: '0px',
                    opacity: '1',
                    height: '150px',
                    width: '150px'
                }, 500);
            });
        });
    </script>
</body>
</html>""",
        "output": "User Interaction:\n1. The page loads with a blue square box (150x150px) on the left side.\n2. The user clicks the 'Move & Fade Box' button.\n3. Animation Output: The blue box smoothly slides 250px to the right, becomes semi-transparent (50% opacity), and grows to 200x200px over a period of 1 second.\n4. The user clicks 'Reset'.\n5. Animation Output: The box smoothly animates back to its original position, size, and full opacity over 0.5 seconds."
    },
    {
        "id": "8",
        "aim": "fetches data asynchronously using jquery",
        "code": """<!DOCTYPE html>
<html>
<head>
    <title>jQuery AJAX Fetch</title>
    <script src="https://code.jquery.com/jquery-3.6.0.min.js"></script>
</head>
<body>

    <h3>User List from API</h3>
    <button id="fetchBtn">Load Users</button>
    <p id="loading" style="display:none; color:blue;">Loading data...</p>
    
    <ul id="userList"></ul>

    <script>
        $(document).ready(function() {
            $("#fetchBtn").click(function() {
                $("#loading").show(); // Show loading indicator
                $("#fetchBtn").prop("disabled", true);
                
                // Asynchronous AJAX Request
                $.ajax({
                    url: "https://jsonplaceholder.typicode.com/users",
                    type: "GET",
                    success: function(data) {
                        $("#loading").hide();
                        let listItems = "";
                        
                        // Loop through first 5 fetched records
                        for (let i = 0; i < 5; i++) {
                            listItems += "<li><strong>" + data[i].name + "</strong> (" + data[i].email + ")</li>";
                        }
                        
                        $("#userList").html(listItems);
                    },
                    error: function(xhr, status, error) {
                        $("#loading").hide();
                        alert("An error occurred: " + error);
                        $("#fetchBtn").prop("disabled", false);
                    }
                });
            });
        });
    </script>
</body>
</html>""",
        "output": "Before Click: Empty list with a 'Load Users' button.\n\nDuring Fetch: The text 'Loading data...' appears temporarily in blue.\n\nAfter Fetch Completion (Output on screen):\n- Leanne Graham (Sincere@april.biz)\n- Ervin Howell (Shanna@melissa.tv)\n- Clementine Bauch (Nathan@yesenia.net)\n- Patricia Lebsack (Julianne.OConner@kory.org)\n- Chelsey Dietrich (Lucio_Hettinger@annie.ca)"
    },
    {
        "id": "9",
        "aim": "php scripts that perform sorting searching and merging",
        "code": """<?php
// PHP Script demonstrating Array Merging, Sorting, and Searching

// 1. Array Definitions
$array1 = array(15, 42, 8, 23);
$array2 = array(10, 5, 99);

echo "<b>Original Array 1:</b> " . implode(", ", $array1) . "<br>";
echo "<b>Original Array 2:</b> " . implode(", ", $array2) . "<br><br>";

// 2. Merging Arrays
$merged_array = array_merge($array1, $array2);
echo "<b>Merged Array:</b> " . implode(", ", $merged_array) . "<br><br>";

// 3. Sorting Array (Ascending Order)
sort($merged_array);
echo "<b>Sorted Merged Array (Ascending):</b> " . implode(", ", $merged_array) . "<br><br>";

// 4. Searching Array
$search_value = 23;
// array_search returns the index if found, or false if not found
$index = array_search($search_value, $merged_array);

if ($index !== false) {
    echo "<span style='color:green;'><b>Search Result:</b> Value {$search_value} was found at index {$index} (0-based) in the sorted array.</span><br>";
} else {
    echo "<span style='color:red;'><b>Search Result:</b> Value {$search_value} was not found.</span><br>";
}

// 5. Searching for a non-existent value
$search_missing = 100;
if (in_array($search_missing, $merged_array)) {
    echo "Value {$search_missing} is present.<br>";
} else {
    echo "<span style='color:red;'><b>Search Result:</b> Value {$search_missing} is NOT present in the array.</span><br>";
}
?>""",
        "output": """Original Array 1: 15, 42, 8, 23
Original Array 2: 10, 5, 99

Merged Array: 15, 42, 8, 23, 10, 5, 99

Sorted Merged Array (Ascending): 5, 8, 10, 15, 23, 42, 99

Search Result: Value 23 was found at index 4 (0-based) in the sorted array.
Search Result: Value 100 is NOT present in the array."""
    },
    {
        "id": "10",
        "aim": "laravel based crud using route controller include the user friendly error manage and bootstrap styling",
        "code": """// 1. web.php (Routes definition)
use App\\Http\\Controllers\\ProductController;

Route::resource('products', ProductController::class);


// 2. ProductController.php (Controller with logic & validation)
namespace App\\Http\\Controllers;
use App\\Models\\Product;
use Illuminate\\Http\\Request;

class ProductController extends Controller
{
    public function index() {
        $products = Product::all();
        return view('products.index', compact('products'));
    }

    public function create() {
        return view('products.create');
    }

    public function store(Request $request) {
        // User-friendly Error Management via validation rules
        $request->validate([
            'name' => 'required|max:255',
            'price' => 'required|numeric|min:0',
        ], [
            'name.required' => 'Please provide a product name.',
            'price.numeric' => 'The price must be a valid number.'
        ]);

        Product::create($request->all());
        return redirect()->route('products.index')
                         ->with('success', 'Product created successfully.');
    }
}


// 3. create.blade.php (View with Bootstrap Styling & Error Display)
@extends('layouts.app')

@section('content')
<div class="container mt-5">
    <h2>Create New Product</h2>
    
    <!-- Displaying Global Errors -->
    @if ($errors->any())
        <div class="alert alert-danger">
            <ul>
                @foreach ($errors->all() as $error)
                    <li>{{ $error }}</li>
                @endforeach
            </ul>
        </div>
    @endif

    <form action="{{ route('products.store') }}" method="POST">
        @csrf
        <div class="mb-3">
            <label class="form-label">Product Name</label>
            <input type="text" name="name" class="form-control @error('name') is-invalid @enderror" value="{{ old('name') }}">
            <!-- Specific Field Error -->
            @error('name')
                <div class="invalid-feedback">{{ $message }}</div>
            @enderror
        </div>

        <div class="mb-3">
            <label class="form-label">Price</label>
            <input type="text" name="price" class="form-control @error('price') is-invalid @enderror" value="{{ old('price') }}">
            @error('price')
                <div class="invalid-feedback">{{ $message }}</div>
            @enderror
        </div>

        <button type="submit" class="btn btn-primary">Save Product</button>
    </form>
</div>
@endsection""",
        "output": """[Scenario 1: Validation Error Output]
If the user submits the form empty, Laravel automatically redirects back. 
The screen displays a red alert box at the top:
- Please provide a product name.
- The price field is required.

Additionally, due to Bootstrap's 'is-invalid' class, the input fields are highlighted with a red border, and the specific error message appears in red text below each respective input field.

[Scenario 2: Successful Submission]
If the user inputs Name="Laptop" and Price="999.99" and clicks Save.
The screen redirects to the Product List page.
A green Bootstrap alert block (.alert-success) appears at the top saying:
"Product created successfully." """
    }
]

os.makedirs('exp_output', exist_ok=True)

for exp in experiments:
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Exp - {exp['id']}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Aim
    doc.add_heading("Aim", level=2)
    doc.add_paragraph(exp['aim'])
    
    # Code
    doc.add_heading("Code", level=2)
    code_paragraph = doc.add_paragraph()
    code_run = code_paragraph.add_run(exp['code'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    # Output
    doc.add_heading("Output", level=2)
    output_paragraph = doc.add_paragraph()
    output_run = output_paragraph.add_run(exp['output'])
    output_run.font.name = 'Consolas'
    output_run.font.size = Pt(10)
    
    filename = f"exp_output/Experiment_{exp['id']}.docx"
    doc.save(filename)
    print(f"Generated {filename}")

# Also generate a single combined document
combined_doc = Document()
combined_doc.add_heading("All Experiments", level=1)

for idx, exp in enumerate(experiments):
    if idx > 0:
        combined_doc.add_page_break()
        
    title = combined_doc.add_heading(f"Exp - {exp['id']}", level=2)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    combined_doc.add_heading("Aim", level=3)
    combined_doc.add_paragraph(exp['aim'])
    
    combined_doc.add_heading("Code", level=3)
    code_paragraph = combined_doc.add_paragraph()
    code_run = code_paragraph.add_run(exp['code'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    combined_doc.add_heading("Output", level=3)
    output_paragraph = combined_doc.add_paragraph()
    output_run = output_paragraph.add_run(exp['output'])
    output_run.font.name = 'Consolas'
    output_run.font.size = Pt(10)

combined_doc.save("exp_output/All_Experiments_Combined.docx")
print("Generated exp_output/All_Experiments_Combined.docx")
