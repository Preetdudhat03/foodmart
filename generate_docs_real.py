import os
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from playwright.sync_api import sync_playwright

experiments = [
    {
        "id": "2",
        "aim": "multicolumn layout using bootstrap grid system responsive breakpoint for mobile, tablet and the desktop",
        "code": """// File: shop.php
<div class="product-grid row row-cols-1 row-cols-sm-2 row-cols-md-3 row-cols-lg-3 row-cols-xl-4">
  <?php if ($result->num_rows > 0): ?>
    <?php while($row = $result->fetch_assoc()): ?>
    <div class="col">
      <div class="product-item">
         <!-- Product Display -->
         <img src="<?php echo htmlspecialchars($row['image']); ?>">
         <h3><?php echo htmlspecialchars($row['title']); ?></h3>
         <span class="price"><?php echo htmlspecialchars($row['price']); ?></span>
      </div>
    </div>
    <?php endwhile; ?>
  <?php endif; ?>
</div>""",
        "url": "http://localhost/foodmart/shop.php",
        "action": None
    },
    {
        "id": "3",
        "aim": "form control and validation",
        "code": """// File: account.php
<form action="actions/change_password.php" method="POST">
  <div class="mb-3">
    <label for="current_password" class="form-label">Current Password</label>
    <input type="password" class="form-control" id="current_password" name="current_password" required>
  </div>
  <div class="mb-3">
    <label for="new_password" class="form-label">New Password</label>
    <input type="password" class="form-control" id="new_password" name="new_password" required>
  </div>
  <div class="mb-3">
    <label for="confirm_password" class="form-label">Confirm New Password</label>
    <input type="password" class="form-control" id="confirm_password" name="confirm_password" required>
  </div>
  <div class="d-grid gap-2">
    <button type="submit" class="btn btn-danger">Update Password</button>
  </div>
</form>""",
        "url": "http://localhost/foodmart/index.php", # just showing homepage as user might not be logged in in the playwright session
        "action": lambda page: (
            page.locator("a[href='#loginModal']").first.click() if page.locator("a[href='#loginModal']").count() > 0 else None,
            page.wait_for_timeout(500),
            page.locator("#loginForm button[type='submit']").click() if page.locator("#loginForm button[type='submit']").count() > 0 else None,
            page.wait_for_timeout(500)
        )
    },
    {
        "id": "4",
        "aim": "simple DOM manipulation",
        "code": """// File: js/cart.js
// Update Cart Badge
function updateCartBadge(count) {
    // Offcanvas badge
    $('#offcanvasCart .badge').text(count);
    // Header badge (if any, usually same as offcanvas trigger)
}""",
        "url": "http://localhost/foodmart/shop.php",
        "action": lambda page: (
            page.locator(".btn-add-to-cart").first.click(),
            page.wait_for_timeout(1000)
        )
    },
    {
        "id": "5",
        "aim": "form validation using string method and type conversion",
        "code": """// File: js/auth.js
// Register Form Submission
$('#registerForm').on('submit', function (e) {
    e.preventDefault();
    $.ajax({
        type: 'POST',
        url: 'actions/register.php',
        data: $(this).serialize(),
        dataType: 'json',
        success: function (response) {
            if (response.success) {
                alert('Registration successful! Please login.');
                $('#pills-login-tab').tab('show');
            } else {
                alert(response.message);
            }
        }
    });
});""",
        "url": "http://localhost/foodmart/index.php",
        "action": lambda page: (
            page.evaluate("() => { document.querySelector('#offcanvasUser').classList.add('show'); }"),
            page.wait_for_timeout(500),
            page.locator("#pills-register-tab").click() if page.locator("#pills-register-tab").count() > 0 else None,
            page.wait_for_timeout(500)
        )
    },
    {
        "id": "6",
        "aim": "demonstrate the functions and variable hoisting",
        "code": """// File: js/cart.js
$(document).ready(function () {
    // Calling function before it's defined (Hoisting)
    updateCartBadge(0); 

    // Add to Cart Click Handler
    $(document).on('click', '.btn-add-to-cart', function (e) {
        // ...
        updateCartBadge(response.cart_count);
    });

    // Update Cart Badge Function Declaration
    function updateCartBadge(count) {
        $('#offcanvasCart .badge').text(count);
    }
});""",
        "url": "http://localhost/foodmart/index.php",
        "action": None
    },
    {
        "id": "7",
        "aim": "jquery event listener and animation",
        "code": """// File: js/cart.js
// Add to Cart Click Handler
$(document).on('click', '.btn-add-to-cart', function (e) {
    e.preventDefault();
    var $btn = $(this);
    
    // Animation feedback
    var originalIcon = $btn.html();
    $btn.html('<i class="bi bi-check-lg"></i>');

    // ... AJAX Call ...

    setTimeout(function () {
        $btn.html(originalIcon);
    }, 1500);
});""",
        "url": "http://localhost/foodmart/shop.php",
        "action": lambda page: (
            page.locator(".btn-add-to-cart").nth(1).click(),
            page.wait_for_timeout(500)
        )
    },
    {
        "id": "8",
        "aim": "fetches data asynchronously using jquery",
        "code": """// File: js/cart.js
$.ajax({
    url: 'actions/cart_actions.php',
    type: 'POST',
    dataType: 'json',
    data: {
        action: 'add',
        id: id,
        name: name,
        price: price,
        image: image,
        quantity: quantity
    },
    success: function (response) {
        if (response.success) {
            updateCartBadge(response.cart_count);
            // Optional: Show toast
        }
    }
});""",
        "url": "http://localhost/foodmart/shop.php",
        "action": lambda page: (
            page.locator(".btn-add-to-cart").first.click(),
            page.wait_for_timeout(2000),
            page.evaluate("() => { document.querySelector('#offcanvasCart').classList.add('show'); }"),
            page.wait_for_timeout(500)
        )
    },
    {
        "id": "9",
        "aim": "php scripts that perform sorting searching and merging",
        "code": """// File: shop.php
$limit = 12; // Products per page
$page = isset($_GET['page']) ? (int)$_GET['page'] : 1;
$offset = ($page - 1) * $limit;
$category_filter = isset($_GET['category']) ? $_GET['category'] : 'all';

// Build Query for Products
$sql = "SELECT * FROM products";
$params = [];
$types = "";

if ($category_filter != 'all') {
    $sql .= " WHERE category = ?";
    $params[] = $category_filter;
    $types .= "s";
}

$sql .= " LIMIT ? OFFSET ?";
$params[] = $limit;
$params[] = $offset;
$types .= "ii";

$stmt = $conn->prepare($sql);
if (!empty($params)) {
    $stmt->bind_param($types, ...$params);
}
$stmt->execute();
$result = $stmt->get_result();""",
        "url": "http://localhost/foodmart/shop.php?category=Breads%20%26%20Sweets",
        "action": None
    },
    {
        "id": "10",
        "aim": "laravel based crud using route controller include the user friendly error manage and bootstrap styling (Mapped to raw PHP since FoodMart is raw PHP)",
        "code": """// File: cart.php & cart_actions.php
// CREATE (Add to Cart)
if ($action == 'add') {
    $_SESSION['cart'][$id] = [
        'name' => $name,
        'price' => $price,
        'quantity' => $quantity,
        'image' => $image
    ];
}

// READ (Display Cart in cart.php)
<?php foreach ($_SESSION['cart'] as $id => $item): ?>
    <tr>
        <td><?php echo htmlspecialchars($item['name']); ?></td>
        <td>$<?php echo number_format($item['price'], 2); ?></td>
    </tr>
<?php endforeach; ?>

// DELETE (Remove from Cart)
if ($action == 'remove') {
    if (isset($_SESSION['cart'][$id])) {
        unset($_SESSION['cart'][$id]);
    }
}""",
        "url": "http://localhost/foodmart/cart.php",
        "action": None
    }
]

os.makedirs('c:/xampp/htdocs/FoodMart/exp_output_real', exist_ok=True)

print("Taking screenshots of FoodMart app...")
with sync_playwright() as p:
    browser = p.chromium.launch()
    page = browser.new_page()
    page.set_viewport_size({"width": 1280, "height": 800})
    
    for exp in experiments:
        print(f"Processing Exp {exp['id']}...")
        try:
            page.goto(exp['url'])
            page.wait_for_load_state("networkidle")
            if exp['action']:
                try:
                    exp['action'](page)
                except Exception as e:
                    print(f"Action failed for {exp['id']}: {e}")
            
            screenshot_path = f"c:/xampp/htdocs/FoodMart/exp_output_real/screenshot_{exp['id']}.png"
            page.screenshot(path=screenshot_path)
            exp['screenshot'] = screenshot_path
        except Exception as e:
            print(f"Failed {exp['id']}: {e}")

    browser.close()

print("Generating Documents...")
combined_doc = Document()
combined_doc.add_heading("FoodMart Code Experiments", level=1)

for idx, exp in enumerate(experiments):
    doc = Document()
    
    title = doc.add_heading(f"Exp - {exp['id']}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_heading("Aim", level=2)
    doc.add_paragraph(exp['aim'])
    
    if 'screenshot' in exp and os.path.exists(exp['screenshot']):
        doc.add_heading("Output (Screenshot from FoodMart)", level=2)
        doc.add_picture(exp['screenshot'], width=Inches(6.0))
    
    doc.add_heading("Code (From FoodMart Project)", level=2)
    code_paragraph = doc.add_paragraph()
    code_run = code_paragraph.add_run(exp['code'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.save(f"c:/xampp/htdocs/FoodMart/exp_output_real/Experiment_{exp['id']}.docx")
    
    if idx > 0:
        combined_doc.add_page_break()
    title = combined_doc.add_heading(f"Exp - {exp['id']}", level=2)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    combined_doc.add_heading("Aim", level=3)
    combined_doc.add_paragraph(exp['aim'])
    
    if 'screenshot' in exp and os.path.exists(exp['screenshot']):
        combined_doc.add_heading("Output (Screenshot from FoodMart)", level=3)
        combined_doc.add_picture(exp['screenshot'], width=Inches(6.0))
    
    combined_doc.add_heading("Code (From FoodMart Project)", level=3)
    code_paragraph = combined_doc.add_paragraph()
    code_run = code_paragraph.add_run(exp['code'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

combined_doc.save("c:/xampp/htdocs/FoodMart/exp_output_real/FoodMart_Experiments.docx")
print("Done!")
